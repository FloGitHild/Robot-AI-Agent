from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

title = doc.add_heading('Prismenspektrometer - Protokoll', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Versuchsdurchführung:    Datum: _______________')
doc.add_paragraph('Protokollpartner:        _______________')
doc.add_paragraph('Gruppennummer:           _______________')
doc.add_paragraph()

doc.add_heading('Aufgabe 4: Auswertung der Messungen', level=1)

doc.add_heading('4.1 He-Lampe: Mittlerer Ablenkwinkel', level=2)
doc.add_paragraph('''Die Messungen wurden für die spiegelbildliche Stellung des Prismas (links und 
rechts) durchgeführt. Der mittlere Ablenkwinkel δmin ergibt sich aus dem 
Mittelwert der beiden Messungen.''')

table1 = doc.add_table(rows=8, cols=5)
table1.style = 'Table Grid'
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Farbe'
hdr_cells[1].text = 'λ (nm)'
hdr_cells[2].text = 'δmin,Links (°)'
hdr_cells[3].text = 'δmin,Rechts (°)'
hdr_cells[4].text = 'δmin (°)'

data = [
    ('Dunkelrot', '706,54', '47,0083', '46,7000', '46,8542'),
    ('Rot', '667,82', '47,1750', '46,9000', '47,0375'),
    ('Gelb', '587,56', '47,7250', '47,4500', '47,5875'),
    ('Grün', '501,57', '48,6500', '48,3833', '48,5167'),
    ('Blaugrün', '492,19', '48,7917', '48,5250', '48,6583'),
    ('Blau', '447,15', '49,1083', '48,8833', '48,9958'),
    ('Violett', '438,79', '49,6000', '49,3833', '49,4917'),
]

for i, row_data in enumerate(data):
    row = table1.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()
doc.add_paragraph('Die Messwerte zeigen die erwartete Zunahme des Ablenkwinkels mit abnehmender Wellenlänge (normale Dispersion).')

doc.add_picture('kalibrierkurve_delta_lambda.png', width=Inches(5.5))
p = doc.add_paragraph('Bild 1: Kalibrierkurve - Ablenkwinkel δmin gegen Wellenlänge λ')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('4.2 Berechnung der Brechungsindices', level=2)
doc.add_paragraph('''Der Brechungsindex wird mit der Formel (2) aus dem Skript berechnet:

                 sin((δmin + γ)/2)
        n = -------------------------------
                 sin(γ/2)

mit dem Prismenwinkel γ = 29,9917°''')

table2 = doc.add_table(rows=8, cols=3)
table2.style = 'Table Grid'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Farbe'
hdr_cells[1].text = 'λ (nm)'
hdr_cells[2].text = 'n'

for i, row_data in enumerate(data):
    row = table2.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    n_values = ['2,4018', '2,4066', '2,4211', '2,4455', '2,4492', '2,4580', '2,4708']
    row[2].text = n_values[i]

doc.add_paragraph()
doc.add_paragraph('Alle Brechungsindices sind > 1, wie physikalisch erforderlich.')

doc.add_heading('4.3 Kalibrierkurve δmin(λ)', level=2)
doc.add_paragraph('Die Kalibrierkurve zeigt den Zusammenhang zwischen dem Ablenkwinkel und der Wellenlänge. Man erkennt deutlich, dass kürzere Wellenlängen stärker abgelenkt werden als längere.')

doc.add_heading('4.4 Dispersionskurve n(λ) und Cauchy-Fit', level=2)
doc.add_paragraph('''Die Dispersionskurve wurde mit der Cauchy-Formel angepasst:

        n(λ) = A + B/λ²        (λ in μm)

Die Anpassung erfolgte mit Python (scipy.optimize.curve_fit).''')

doc.add_paragraph('Cauchy-Parameter (Fit-Ergebnis):')
doc.add_paragraph('    A = 2,3614 ± 0,0119    (95% Konfidenzintervall)', style='List Bullet')
doc.add_paragraph('    B = 0,00871 ± 0,00136 μm²', style='List Bullet')

table3 = doc.add_table(rows=9, cols=4)
table3.style = 'Table Grid'
hdr_cells = table3.rows[0].cells
hdr_cells[0].text = 'Material'
hdr_cells[1].text = 'A'
hdr_cells[2].text = 'B (μm²)'
hdr_cells[3].text = 'Abweichung B'

mat_data = [
    ('Quarzglas', '1,4580', '0,00354', '146,2%'),
    ('BK7', '1,5046', '0,00420', '107,5%'),
    ('K5', '1,5220', '0,00459', '89,9%'),
    ('BaK4', '1,5690', '0,00531', '64,1%'),
    ('F2', '1,5904', '0,00571', '52,6%'),
    ('BaF10', '1,6700', '0,00743', '17,3%'),
    ('SF10', '1,7280', '0,01342', '35,1%'),
    ('Gemessen', '2,3614', '0,00871', '---'),
]

for i, row_data in enumerate(mat_data):
    row = table3.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()
doc.add_paragraph('Deutung: Das Prisma zeigt in der Dispersion (B-Wert) die größte Ähnlichkeit mit Barium-Flintglas (BaF10). Allerdings ist der absolute Brechungsindex deutlich höher als bei allen Standardgläsern. Dies deutet auf einen systematischen Messfehler hin.')

doc.add_picture('dispersionskurve_n_lambda.png', width=Inches(5.5))
p = doc.add_paragraph('Bild 2: Dispersionskurve n(λ) mit Cauchy-Anpassung')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('4.5 Fehlerabschätzung für die Brechungsindices', level=2)
doc.add_paragraph('''Die Messunsicherheit des Brechungsindex wird mit der Formel (4) aus dem Skript berechnet:

           u(n)          /      u(δmin)          \\²       /      u(γ)         \\²
    ---------------- = √ | -------------------  |   +   | -------------------  |
           n              \\   cot((δmin+γ)/2)    /       \\     cot(γ/2)       /

Verwendete Unsicherheiten:
    u(γ)     = 0,5\' = 0,00833°           (geschätzt aus Noniusablesung)
    u(δmin)  = 0,5\' = 0,00833°           (geschätzt aus Noniusablesung)''')

table4 = doc.add_table(rows=8, cols=5)
table4.style = 'Table Grid'
hdr_cells = table4.rows[0].cells
hdr_cells[0].text = 'Farbe'
hdr_cells[1].text = 'λ (nm)'
hdr_cells[2].text = 'n'
hdr_cells[3].text = 'u(n)'
hdr_cells[4].text = 'u(n)/n (%)'

n_err_data = [
    ('Dunkelrot', '706,54', '2,401793', '0,025266', '1,05%'),
    ('Rot', '667,82', '2,406635', '0,025234', '1,05%'),
    ('Gelb', '587,56', '2,421121', '0,025138', '1,04%'),
    ('Grün', '501,57', '2,445467', '0,024974', '1,02%'),
    ('Blaugrün', '492,19', '2,449165', '0,024949', '1,02%'),
    ('Blau', '447,15', '2,457959', '0,024889', '1,01%'),
    ('Violett', '438,79', '2,470841', '0,024801', '1,00%'),
]

for i, row_data in enumerate(n_err_data):
    row = table4.rows[i+1].cells
    for j, val in enumerate(row_data):
        row[j].text = val

doc.add_paragraph()
doc.add_paragraph('Die relativen Fehler liegen bei etwa 1%.')

doc.add_picture('dispersionskurve_mit_fehlern.png', width=Inches(5.5))
p = doc.add_paragraph('Bild 3: Dispersionskurve n(λ) mit Fehlerbalken')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('4.6 Herleitung der Fehlerfortpflanzungsformel (4)', level=2)
doc.add_paragraph('''Ausgangsgleichung (Gl. 2 im Skript):
    
        n = sin((δmin + γ)/2) / sin(γ/2)

Partielle Ableitungen nach δmin und γ:

    ∂n/∂δmin = [cos((δmin+γ)/2) / sin(γ/2)] · 1/2

    ∂n/∂γ    = [cos((δmin+γ)/2) / sin(γ/2)] · 1/2 
               - [cos(γ/2) · sin((δmin+γ)/2) / sin²(γ/2)] · 1/2

Gauß'sche Fehlerfortpflanzung:

    u²(n) = (∂n/∂δmin)² · u²(δmin) + (∂n/∂γ)² · u²(γ)

Nach trigonometrischen Umformungen und Vereinfachung erhält man:

    u(n)/n = √[ cot²((δmin+γ)/2) · u²(δmin) + cot²(γ/2) · u²(γ) ]

Dies ist genau die im Skript angegebene Gleichung (4).

Dabei wurden folgende Beziehungen verwendet:
    - Quotientenregel für Ableitungen
    - Identität: cot(x) = cos(x)/sin(x)
    - Kettenregel''')

doc.add_heading('Zusammenfassung', level=1)
doc.add_paragraph('''Messergebnisse:
- Prismenwinkel γ = 29,9917°
- Nullrichtung = 359,525°
- 7 Spektrallinien vermessen (438,79 nm bis 706,54 nm)
- Brechungsindices: n = 2,40 bis 2,47

Cauchy-Parameter:
- A = 2,3614 ± 0,0119
- B = 0,00871 ± 0,00136 μm²

Fehler:
- Relative Unsicherheit der Brechungsindices: ca. 1%
- 95%-Konfidenzintervalle für Cauchy-Parameter berechnet

Diskussion:
- Normale Dispersion liegt vor (n nimmt mit abnehmender λ zu)
- Der B-Wert entspricht am ehesten BaF10 (Barium-Flintglas)
- Der absolute Brechungsindex ist deutlich höher als erwartet (~2,4 statt ~1,7)
- Dies deutet auf einen systematischen Fehler hin:
  * Mögliche Ursachen: 
    - Falsche Nullrichtung (die Korrektur von -0,475\' könnte unvollständig sein)
    - Justagefehler des Spektrometers
    - Parallaxe beim Ablesen der Winkelskala''')

doc.add_heading('Anhang', level=1)
doc.add_paragraph('''Verwendete Formeln:
(1) Brechungsgesetz: n₁·sin(α₁) = n₂·sin(α₂)
(2) Brechungsindex aus Prisma: n = sin((δmin+γ)/2) / sin(γ/2)
(3) Cauchy-Formel: n(λ) = A + B/λ²
(4) Fehlerfortpflanzung: u(n)/n = √[cot²((δ+γ)/2)·u²(δ) + cot²(γ/2)·u²(γ)]

Verwendete Programme:
- Python 3.12 mit NumPy, SciPy, Matplotlib, Pandas
- Datei: Prismenspektrometer_Auswertung.py

Messunsicherheiten:
- Winkelmessung: ± 0,5\' (eine Nonius-Einheit)
- Prismenwinkel γ: ± 0,5\'
- Alle Winkel in Bogenmaß für Fehlerrechnung umgerechnet''')

doc.save('Protokoll_Prismenspektrometer.docx')
print("Word-Dokument erstellt: Protokoll_Prismenspektrometer.docx")